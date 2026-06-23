from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT / "downloads"


def build_pdf() -> None:
    DOWNLOADS.mkdir(exist_ok=True)
    path = DOWNLOADS / "creativex2026-knowledge-brief.pdf"
    styles = getSampleStyleSheet()
    story = [
        Paragraph("CreativEX2026 Knowledge Brief", styles["Title"]),
        Spacer(1, 14),
        Paragraph(
            "A concise executive brief for the CreativEX2026 Knowledge Platform. "
            "The platform connects knowledge articles, speaker pages, search, downloads, "
            "member activity, portfolio creation, and impact evaluation.",
            styles["BodyText"],
        ),
        Spacer(1, 14),
        Paragraph("Core Platform Modules", styles["Heading2"]),
        Table(
            [
                ["Module", "Purpose"],
                ["Knowledge Library", "Collect and search learning resources."],
                ["Speaker Pages", "Capture structured knowledge from expert sessions."],
                ["Dashboard", "Track articles, downloads, visitors, and search queries."],
                ["Impact Framework", "Connect learning activity to measurable outcomes."],
            ],
            colWidths=[170, 330],
        ),
    ]
    story[-1].setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#116466")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d9dee8")),
                ("PADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    SimpleDocTemplate(str(path), pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54).build(story)


def build_docx() -> None:
    DOWNLOADS.mkdir(exist_ok=True)
    path = DOWNLOADS / "creativex2026-impact-worksheet.docx"
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    doc.add_heading("CreativEX2026 Impact Evaluation Worksheet", level=0)
    doc.add_paragraph("Use this worksheet to define indicators, collect evidence, and report learning impact.")
    rows = [
        ("Inputs", "Budget, partners, speakers, data, platform tools", "What resources are committed?"),
        ("Activities", "Learning, search, upload, portfolio, review", "What actions happened?"),
        ("Outputs", "Articles, downloads, active users, submitted work", "What was produced?"),
        ("Outcomes", "Skills, collaboration, quality of work", "What changed for participants?"),
        ("Impact", "Creative economy value, city brand, policy decisions", "What longer-term value emerged?"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Layer"
    header[1].text = "Evidence"
    header[2].text = "Question"
    for layer, evidence, question in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = evidence
        cells[2].text = question
    doc.add_paragraph("Next step: assign owners, data sources, and review cadence for each indicator.")
    doc.save(path)


if __name__ == "__main__":
    build_pdf()
    build_docx()
