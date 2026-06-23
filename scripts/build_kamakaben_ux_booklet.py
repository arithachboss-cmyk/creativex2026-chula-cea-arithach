from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT / "downloads"
DOCX_PATH = DOWNLOADS / "kamakaben-ux-research-booklet.docx"
MD_PATH = ROOT / "kamakaben-ux-research-booklet.md"


TITLE = "Kamakaben UX Research Booklet"
SUBTITLE = "เม็ดต่อของประสบการณ์ผู้ใช้: กลิ่นไทย ความจำ และการเรียนรู้เชิงวัฒนธรรม"
FONT = "Thonburi"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(9.5)
    run.bold = bold


def add_paragraph(doc, text, style=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:cs"), FONT)
    r.font.size = Pt(10.5)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r._element.rPr.rFonts.set(qn("w:cs"), FONT)
        r.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r._element.rPr.rFonts.set(qn("w:cs"), FONT)
        r.font.size = Pt(10)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run._element.rPr.rFonts.set(qn("w:cs"), FONT)
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def build_markdown():
    md = f"""# {TITLE}

## {SUBTITLE}

เวอร์ชันนี้คือเอกสารวิชาประกอบสำหรับใช้เป็น “เม็ดต่อ” ของ UX: เปลี่ยนภาพอ้างอิงการทำกลิ่นและการทดสอบความรู้สึกให้กลายเป็นระบบประสบการณ์ผู้ใช้ของ Kamakaben

## Thesis

Kamakaben ไม่ควรเป็นแค่แบรนด์น้ำหอม แต่ควรเป็นแพลตฟอร์มเรียนรู้กลิ่นไทยที่ให้ผู้ใช้เลือกภาพ ความทรงจำ สถานที่ และอารมณ์ แล้วแปลงเป็นเส้นทางกลิ่นที่บันทึกได้ แชร์ได้ และต่อยอดเป็นวัฒนธรรมร่วมสมัยได้

## Research Inputs

- ภาพอ้างอิงจากผู้ใช้: กระบวนการผสมกลิ่น 50ml, การเลือกสูตร, การจดคะแนน, การจับคู่ภาพกับกลิ่น, การอ่านปฏิกิริยาสมอง/ความรู้สึก, และแบบสอบถามหลังดม
- Benchmark: Karmakamet สะกดเป็น `Karmakamet`; ใช้เป็นกรณีศึกษาเรื่อง scent, memory, craftsmanship, heritage, store experience โดยไม่ลอกอัตลักษณ์
- Academic grounding: odor-evoked memory, sensory/affective UX, scent marketing, participatory learning, และ cultural experience design

## UX Seed

1. Scent Image Prompt: ให้ผู้ใช้เลือกภาพ/สถานที่/ความทรงจำก่อนเลือกกลิ่น
2. 50ml Lab Ritual: ทำให้ขั้นตอนทดลองกลิ่นกลายเป็นพิธีกรรมเรียนรู้ ไม่ใช่แค่การซื้อสินค้า
3. Memory Fit Score: ให้คะแนนว่ากลิ่นตรงกับภาพ ความรู้สึก และเรื่องเล่าของผู้ใช้แค่ไหน
4. Scent Passport: บันทึกกลิ่น สถานที่ อารมณ์ และบทเรียนทางวัฒนธรรม
5. Learning Dashboard: รวมคะแนนการเรียนรู้ ดาวน์โหลด และร่องรอยการค้นหา
6. Global Thai Scent Map: แผนที่กลิ่นไทยที่คนทั่วโลกเพิ่มประสบการณ์ของตนเองได้

## Prototype Flow

เลือกภาพ → เลือก territory กลิ่นไทย → ผสมสูตรทดลอง 50ml → ดมและให้คะแนน → เขียน memory note → รับ scent passport entry → แชร์/ดาวน์โหลด/นำไปต่อยอดเวิร์กช็อป

## Evidence Notes

- กลิ่นสัมพันธ์กับความจำและอารมณ์ จึงเหมาะกับ UX ที่ใช้ personal memory และ cultural learning
- Karmakamet วางตำแหน่งกลิ่นกับ memory, inner world, craftsmanship และ heritage ได้แข็งแรง เป็น benchmark ฝั่ง Thai fragrance culture
- ภาพอ้างอิงเกาหลีชี้ให้เห็นว่าประสบการณ์กลิ่นที่ดีควรมีทั้ง lab ritual, formula sheet, scoring, sensory testing, และ questionnaire

## Next Build

- ทำหน้า Scent Lab ในเว็บไซต์
- เพิ่มแบบฟอร์ม Memory Fit Score ใน Scent Passport
- ทำ resource download สำหรับ workshop facilitator
- สร้างชุด QR production จริงสำหรับ scent territory cards
- เก็บข้อมูล UX pilot 20 คนแรก

## Sources

- Karmakamet official website: https://karmakamet.co.th/en/
- Karmakamet About: https://karmakamet.co.th/about/
- Karmakamet Craftsmanship: https://karmakamet.co.th/en/craftsmanship/
- Herz, R. S. (2016). The role of odor-evoked memory in psychological and physiological health. Brain Sciences. https://doi.org/10.3390/brainsci6030022
- Larsson et al. (2014). Olfactory LOVER. Frontiers in Psychology. https://doi.org/10.3389/fpsyg.2014.00312
- Sugiyama et al. (2015). Proustian products are preferred. Chemosensory Perception. https://doi.org/10.1007/s12078-015-9182-y
"""
    MD_PATH.write_text(md, encoding="utf-8")


def build_docx():
    DOWNLOADS.mkdir(exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), FONT)
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(TITLE)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:cs"), FONT)
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(SUBTITLE)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:cs"), FONT)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(82, 82, 82)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Academic companion booklet | UX seed document | June 2026")
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:cs"), FONT)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(doc, "1. บทตั้งต้น", 1)
    add_paragraph(
        doc,
        "เล่มนี้สรุปฐานคิดวิชาการและเม็ดต่อของประสบการณ์ผู้ใช้สำหรับ Kamakaben: "
        "แบรนด์กลิ่นไทยที่ทำให้คนทั่วโลกเรียนรู้ตัวตนของประเทศไทยผ่านกลิ่น ความทรงจำ และพิธีกรรมการเลือกกลิ่นของตนเอง",
    )
    add_paragraph(
        doc,
        "โจทย์ UX ไม่ใช่การทำหน้าน้ำหอมให้สวยเท่านั้น แต่คือการออกแบบเส้นทางที่ผู้ใช้รู้สึกว่า "
        "กลิ่นหนึ่งกลิ่นพาเขาเข้าใกล้ความทรงจำ สถานที่ วัฒนธรรม และตัวตนได้จริง",
        bold=True,
        color="1F3A5F",
    )

    add_heading(doc, "2. สิ่งที่ถอดจากภาพอ้างอิง", 1)
    add_table(
        doc,
        ["องค์ประกอบจากภาพ", "ความหมายเชิง UX", "สิ่งที่ควรทำใน Kamakaben"],
        [
            ["ผสมกลิ่น 50ml", "ทำให้การทดลองมีกติกาและขนาดที่จับต้องได้", "สร้าง 50ml Scent Lab Ritual สำหรับเวิร์กช็อป"],
            ["ขวดสูตร / pipette / เครื่องชั่ง", "ผู้ใช้เห็นความละเอียดและความน่าเชื่อถือ", "โชว์ formula note แบบเข้าใจง่าย ไม่เปิดสูตรลับ"],
            ["สูตรบนเอกสารและรถเข็น sample", "กลิ่นกลายเป็นระบบความรู้ ไม่ใช่ความรู้สึกลอย ๆ", "ทำ scent territory card พร้อมส่วนประกอบวัฒนธรรม"],
            ["ใบคะแนน result", "ประสบการณ์กลิ่นต้องวัดผลได้", "เพิ่ม Memory Fit Score และ Cultural Learning Score"],
            ["เทียบกลิ่นกับภาพที่เลือก", "ภาพเป็นตัวกระตุ้นความทรงจำก่อนกลิ่น", "ให้ผู้ใช้เลือกภาพ/สถานที่/อารมณ์ก่อนเริ่ม"],
            ["อุปกรณ์อ่านปฏิกิริยา/แบบสอบถาม", "วัดทั้งความรู้สึกและร่องรอยการเรียนรู้", "ใช้แบบสอบถาม UX pilot ก่อนใช้เครื่องมือชีวสัญญาณจริง"],
        ],
        widths=[1.65, 2.2, 2.65],
    )

    add_heading(doc, "3. Benchmark: Karmakamet", 1)
    add_paragraph(
        doc,
        "สะกดแบรนด์อ้างอิงให้ถูกต้องเป็น Karmakamet. ใช้เป็น benchmark ด้าน Thai fragrance culture, heritage, craftsmanship และประสบการณ์หน้าร้าน แต่ Kamakaben ต้องมีแกนของตัวเอง: เปิดให้ผู้ใช้ร่วมสร้างความหมายของกลิ่นไทยผ่านข้อมูล ความทรงจำ และการเรียนรู้",
    )
    add_bullets(
        doc,
        [
            "Karmakamet สื่อสารเรื่อง scent as a tool, memory, inner world และ craftsmanship ชัดเจน",
            "มีหมวดสินค้าและ touchpoint หลายแบบ เช่น aromatic oil, essential oil, candle, incense, dry perfume, liquid fragrance และ body care",
            "บทเรียนสำหรับ Kamakaben คือความหนักแน่นของ narrative heritage แต่ต้องเพิ่มชั้น platform, learning dashboard และ participatory UX",
        ],
    )

    add_heading(doc, "4. ฐานวิชาการที่รองรับ", 1)
    add_table(
        doc,
        ["ฐานคิด", "ใจความ", "แปลเป็น UX"],
        [
            ["Odor-evoked memory", "กลิ่นกระตุ้นความทรงจำและอารมณ์ได้แรงกว่าสิ่งเร้าหลายประเภท", "ให้ผู้ใช้บันทึก memory note ทุกครั้งที่ดม"],
            ["Affective design", "ประสบการณ์ที่ดีต้องวัดความรู้สึก ไม่ใช่แค่ conversion", "ใช้ mood before/after และ fit score"],
            ["Scent marketing", "กลิ่นช่วยสร้างการจำแบรนด์และบรรยากาศ", "สร้าง signature scent territories สำหรับช่องทางออนไลน์และออฟไลน์"],
            ["Cultural learning", "การเรียนรู้เกิดเมื่อผู้ใช้เชื่อมประสบการณ์ส่วนตัวกับความหมายทางวัฒนธรรม", "ใส่ story card ของวัตถุดิบ สถานที่ และพิธีกรรมไทย"],
        ],
        widths=[1.65, 2.45, 2.4],
    )

    add_heading(doc, "5. UX Seed Model", 1)
    add_numbered(
        doc,
        [
            "Scent Image Prompt: เริ่มจากภาพ สถานที่ หรือความทรงจำก่อนเลือกกลิ่น",
            "Thai Scent Territory: ให้เลือก territory เช่น ฝนบนไม้สัก มะลิยามค่ำ เช้าวัด ตลาดริมน้ำ ทุ่งทอง",
            "50ml Lab Ritual: ทำสูตรทดลองขนาดเล็กเพื่อให้ผู้ใช้เห็น craft และมี ownership",
            "Memory Fit Score: ให้คะแนน 1-5 ว่ากลิ่นตรงกับภาพ/ความรู้สึก/เรื่องเล่าเพียงใด",
            "Scent Passport: เก็บ entry เป็นประวัติการเรียนรู้ส่วนตัว",
            "Learning Dashboard: รวมบทความ ดาวน์โหลด จำนวนการค้นหา และ insight เพื่อพัฒนากลิ่นรอบต่อไป",
        ],
    )

    add_heading(doc, "6. Prototype Flow", 1)
    add_paragraph(
        doc,
        "เลือกภาพ → เลือก scent territory → รับ story card → ผสม/ทดลอง 50ml → ดมแบบ blind หรือ guided → ให้คะแนน → เขียน memory note → บันทึกลง Scent Passport → แชร์/ดาวน์โหลด certificate หรือ workshop sheet",
        bold=True,
        color="1F3A5F",
    )

    add_heading(doc, "7. Metrics สำหรับ UX Pilot", 1)
    add_table(
        doc,
        ["ตัวชี้วัด", "คำถามที่วัด", "ผลลัพธ์ที่ต้องการ"],
        [
            ["Memory Fit", "กลิ่นนี้ตรงกับภาพ/ความทรงจำที่เลือกแค่ไหน", "เฉลี่ย 4.0/5 ขึ้นไป"],
            ["Cultural Recall", "หลังดม ผู้ใช้จำเรื่องวัฒนธรรมไทยอะไรได้", "ตอบได้อย่างน้อย 1 insight"],
            ["Emotional Shift", "อารมณ์ก่อน-หลังดมเปลี่ยนอย่างไร", "รู้สึกสงบ/สนใจ/เชื่อมโยงมากขึ้น"],
            ["Share Intent", "อยากแชร์ entry หรือชวนคนอื่นมาลองไหม", "มากกว่า 60%"],
            ["Repeat Path", "อยากทดลอง territory อื่นต่อไหม", "มากกว่า 50%"],
        ],
        widths=[1.45, 3.1, 1.95],
    )

    add_heading(doc, "8. ข้อควรระวัง", 1)
    add_bullets(
        doc,
        [
            "ห้ามสื่อสารเป็นการรักษาทางการแพทย์โดยไม่มีหลักฐานและการอนุมัติที่เกี่ยวข้อง ใช้คำว่า emotional reset, reflective ritual หรือ learning experience จะปลอดภัยกว่า",
            "ข้อมูลอารมณ์ ความทรงจำ และชีวสัญญาณเป็นข้อมูลอ่อนไหว ต้องมี consent, privacy notice และสิทธิ์ลบข้อมูล",
            "Benchmark Karmakamet ใช้เพื่อเรียนรู้ตำแหน่งตลาดและภาษาประสบการณ์ ไม่ใช่การคัดลอกชื่อ กลิ่น ภาพ หรือ story world",
        ],
    )

    add_heading(doc, "9. เม็ดต่อที่ต้องทำทันที", 1)
    add_numbered(
        doc,
        [
            "เพิ่มหน้า Kamakaben Scent Lab ในเว็บไซต์",
            "ต่อแบบฟอร์ม Memory Fit Score เข้ากับ Scent Passport",
            "สร้าง workshop worksheet เป็น PDF/DOCX ให้ดาวน์โหลด",
            "ทำ QR matrix จริงสำหรับ scent territory cards",
            "เก็บ UX pilot 20 คนแรก แล้วสรุป insight เป็น dashboard",
        ],
    )

    add_heading(doc, "10. Sources", 1)
    add_bullets(
        doc,
        [
            "Karmakamet official website: https://karmakamet.co.th/en/",
            "Karmakamet About: https://karmakamet.co.th/about/",
            "Karmakamet Craftsmanship: https://karmakamet.co.th/en/craftsmanship/",
            "Herz, R. S. (2016). The role of odor-evoked memory in psychological and physiological health. Brain Sciences. https://doi.org/10.3390/brainsci6030022",
            "Larsson et al. (2014). Olfactory LOVER. Frontiers in Psychology. https://doi.org/10.3389/fpsyg.2014.00312",
            "Sugiyama et al. (2015). Proustian products are preferred. Chemosensory Perception. https://doi.org/10.1007/s12078-015-9182-y",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Kamakaben UX Research Booklet | seed for user experience")
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run._element.rPr.rFonts.set(qn("w:cs"), FONT)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_markdown()
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)
